"""
标准文档转换工具 - 将 Word/PDF 标准文档转换为 JSON 格式

使用场景：
1. 你拿到一份 Word 格式的标准文档（如：《软件测试规范.docx》）
2. 运行此工具，自动提取规则并生成 JSON
3. 可选：使用 LLM 辅助提取（更智能）

转换策略：
- 基础版：规则模板匹配（快速，适合格式规范的文档）
- 增强版：LLM 辅助提取（智能，适合复杂文档）
"""
import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from loguru import logger
from docx import Document

from ..models.document import Standard, Category, Rule, CheckType, Severity


class StandardConverter:
    """标准文档转换器"""
    
    def __init__(self, use_llm: bool = False):
        """
        Args:
            use_llm: 是否使用 LLM 辅助提取（需要配置 LLM）
        """
        self.use_llm = use_llm
        self.llm_client = None
        
        if use_llm:
            # TODO: 初始化 LLM 客户端
            logger.info("🤖 LLM 辅助模式已启用")
    
    def convert_word_to_json(
        self,
        word_path: str,
        output_path: Optional[str] = None,
        protocol_id: Optional[str] = None,
        protocol_name: Optional[str] = None
    ) -> str:
        """
        将 Word 标准文档转换为 JSON
        
        Args:
            word_path: Word 文档路径
            output_path: 输出 JSON 路径（默认：同目录下同名.json）
            protocol_id: 协议ID（默认：从文件名提取）
            protocol_name: 协议名称（默认：从文档标题提取）
        
        Returns:
            输出文件路径
        """
        logger.info(f"📄 开始转换标准文档: {word_path}")
        
        # 解析 Word 文档
        doc = Document(word_path)
        
        # 提取元数据
        if not protocol_id:
            protocol_id = Path(word_path).stem.upper().replace(" ", "_")
        
        if not protocol_name:
            protocol_name = self._extract_title(doc)
        
        logger.info(f"   协议ID: {protocol_id}")
        logger.info(f"   协议名称: {protocol_name}")
        
        # 提取规则
        categories = self._extract_rules(doc)
        
        # 构建标准对象
        standard = Standard(
            protocol_id=protocol_id,
            name=protocol_name,
            version="1.0",
            description=f"从 {Path(word_path).name} 自动提取",
            categories=categories
        )
        
        # 保存为 JSON
        if not output_path:
            output_path = str(Path(word_path).with_suffix('.json'))
        
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(
                standard.model_dump(),
                f,
                ensure_ascii=False,
                indent=2
            )
        
        logger.info(f"✅ 转换完成: {output_path}")
        logger.info(f"   共提取 {len(categories)} 个分类, {sum(len(c.rules) for c in categories)} 条规则")
        
        return output_path
    
    def _extract_title(self, doc: Document) -> str:
        """提取文档标题"""
        for para in doc.paragraphs[:5]:  # 只看前5段
            text = para.text.strip()
            if text and len(text) < 50:
                # 判断是否是标题样式
                if para.style.name in ['Title', 'Heading 1'] or \
                   (para.runs and para.runs[0].bold and para.runs[0].font.size and para.runs[0].font.size.pt > 14):
                    return text
        
        return "未命名标准"
    
    def _extract_rules(self, doc: Document) -> List[Category]:
        """
        提取规则（核心逻辑）
        
        策略：
        1. 识别章节标题（作为 Category）
        2. 提取规则条目（编号 + 描述）
        3. 分析规则类型和关键词
        """
        categories = []
        current_category = None
        rule_counter = 1
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 判断是否是章节标题
            if self._is_section_heading(para, text):
                # 保存上一个分类
                if current_category and current_category["rules"]:
                    categories.append(Category(**current_category))
                
                # 创建新分类
                current_category = {
                    "category": self._clean_section_title(text),
                    "rules": []
                }
                logger.debug(f"   发现章节: {current_category['category']}")
            
            # 判断是否是规则条目
            elif current_category and self._is_rule_item(text):
                rule = self._parse_rule(text, rule_counter)
                if rule:
                    current_category["rules"].append(rule)
                    rule_counter += 1
                    logger.debug(f"      提取规则: {rule.description[:30]}...")
        
        # 保存最后一个分类
        if current_category and current_category["rules"]:
            categories.append(Category(**current_category))
        
        return categories
    
    def _is_section_heading(self, para, text: str) -> bool:
        """判断是否是章节标题"""
        # 方法1：样式判断
        if para.style.name.startswith('Heading'):
            return True
        
        # 方法2：格式判断（加粗 + 字号大）
        if para.runs:
            first_run = para.runs[0]
            if first_run.bold and first_run.font.size and first_run.font.size.pt >= 12:
                return True
        
        # 方法3：文本模式判断
        # 匹配：一、二、三 或 1. 2. 3. 或 第一章、第二节
        patterns = [
            r'^[一二三四五六七八九十]+[、．]',  # 一、
            r'^\d+[\.\．、]',  # 1.
            r'^第[一二三四五六七八九十\d]+[章节条款]',  # 第一章
            r'^[\d\.]+\s+[^\d]',  # 1.1 标题
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                # 排除过长的文本（可能是正文）
                if len(text) < 50:
                    return True
        
        return False
    
    def _clean_section_title(self, text: str) -> str:
        """清理章节标题（去除编号）"""
        # 去除常见的编号格式
        text = re.sub(r'^[一二三四五六七八九十]+[、．]\s*', '', text)
        text = re.sub(r'^\d+[\.\．、]\s*', '', text)
        text = re.sub(r'^第[一二三四五六七八九十\d]+[章节条款]\s*', '', text)
        text = re.sub(r'^[\d\.]+\s+', '', text)
        
        return text.strip()
    
    def _is_rule_item(self, text: str) -> bool:
        """判断是否是规则条目"""
        # 规则条目特征：
        # 1. 有编号（如：1）、（1）、①）
        # 2. 包含"应"、"必须"、"不得"、"应当"等关键词
        # 3. 长度适中（20-200字）
        
        if len(text) < 20 or len(text) > 500:
            return False
        
        # 编号模式
        number_patterns = [
            r'^\d+[）\)、．]',  # 1）、1)、1、、1.
            r'^[（\(]\d+[）\)]',  # (1)、（1）
            r'^[①②③④⑤⑥⑦⑧⑨⑩]',  # ①
        ]
        
        has_number = any(re.match(p, text) for p in number_patterns)
        
        # 规则关键词
        rule_keywords = ['应', '必须', '不得', '应当', '需要', '要求', '禁止', '不应', '宜', '可']
        has_keyword = any(kw in text for kw in rule_keywords)
        
        return has_number or has_keyword
    
    def _parse_rule(self, text: str, rule_id: int) -> Optional[Rule]:
        """解析单条规则"""
        try:
            # 清理编号
            description = self._clean_rule_number(text)
            
            # 分析规则类型
            check_type = self._infer_check_type(description)
            
            # 提取关键词
            keywords = self._extract_keywords(description)
            
            # 判断严重程度
            severity = self._infer_severity(description)
            
            return Rule(
                rule_id=f"R{rule_id:03d}",
                description=description,
                check_type=check_type,
                keywords=keywords,
                positive_examples=[],  # 需要手动补充或 LLM 生成
                negative_examples=[],
                severity=severity
            )
        except Exception as e:
            logger.warning(f"解析规则失败: {text[:50]}... - {e}")
            return None
    
    def _clean_rule_number(self, text: str) -> str:
        """清理规则编号"""
        text = re.sub(r'^\d+[）\)、．]\s*', '', text)
        text = re.sub(r'^[（\(]\d+[）\)]\s*', '', text)
        text = re.sub(r'^[①②③④⑤⑥⑦⑧⑨⑩]\s*', '', text)
        return text.strip()
    
    def _infer_check_type(self, text: str) -> CheckType:
        """推断检查类型"""
        # 格式类关键词
        format_keywords = ['格式', '字体', '字号', '标点', '缩进', '对齐', '页边距', '行距']
        if any(kw in text for kw in format_keywords):
            return CheckType.FORMAT
        
        # 结构类关键词
        structure_keywords = ['结构', '章节', '目录', '顺序', '层次', '组成']
        if any(kw in text for kw in structure_keywords):
            return CheckType.STRUCTURE
        
        # 默认为语义类
        return CheckType.SEMANTIC
    
    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词（简单版）"""
        # 提取名词和动词（简化版，可用 jieba 分词优化）
        keywords = []
        
        # 常见关键词模式
        common_keywords = [
            '标题', '正文', '摘要', '关键词', '目录', '页码', '页眉', '页脚',
            '字体', '字号', '加粗', '斜体', '下划线', '标点', '缩进', '对齐',
            '图表', '表格', '公式', '引用', '参考文献', '附录',
            '准确', '简洁', '完整', '规范', '清晰', '一致'
        ]
        
        for kw in common_keywords:
            if kw in text:
                keywords.append(kw)
        
        return keywords[:5]  # 最多5个
    
    def _infer_severity(self, text: str) -> Severity:
        """推断严重程度"""
        # 高优先级关键词
        high_keywords = ['必须', '禁止', '不得', '严禁']
        if any(kw in text for kw in high_keywords):
            return Severity.HIGH
        
        # 低优先级关键词
        low_keywords = ['宜', '可', '建议', '推荐']
        if any(kw in text for kw in low_keywords):
            return Severity.LOW
        
        return Severity.MEDIUM
    
    def convert_with_llm(self, word_path: str) -> str:
        """
        使用 LLM 辅助转换（更智能）
        
        TODO: 实现 LLM 辅助提取
        - 自动识别规则
        - 生成正反例
        - 优化描述
        """
        raise NotImplementedError("LLM 辅助转换功能待实现")


def main():
    """命令行工具"""
    import argparse
    
    parser = argparse.ArgumentParser(description="标准文档转换工具")
    parser.add_argument("input", help="输入 Word 文档路径")
    parser.add_argument("-o", "--output", help="输出 JSON 路径")
    parser.add_argument("--id", help="协议ID")
    parser.add_argument("--name", help="协议名称")
    parser.add_argument("--llm", action="store_true", help="使用 LLM 辅助")
    
    args = parser.parse_args()
    
    converter = StandardConverter(use_llm=args.llm)
    converter.convert_word_to_json(
        word_path=args.input,
        output_path=args.output,
        protocol_id=args.id,
        protocol_name=args.name
    )


if __name__ == "__main__":
    main()

