"""
文档解析器 - 支持 Word、PDF 等格式
"""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from typing import List, Dict, Any, Optional
from loguru import logger
import re


class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        self.current_section = ""
        self.section_hierarchy = []
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        解析 Word 文档，保留结构信息
        
        Args:
            file_path: 文档路径
        
        Returns:
            结构化的文档数据
        """
        try:
            doc = Document(file_path)
            
            # 提取文档结构
            structure = {
                "title": self._extract_title(doc),
                "sections": [],
                "paragraphs": [],
                "metadata": {
                    "total_paragraphs": 0,
                    "total_chars": 0
                }
            }
            
            current_section = None
            paragraph_index = 0
            
            for element in doc.element.body:
                # 处理段落
                if isinstance(element, CT_P):
                    para = Paragraph(element, doc)
                    text = para.text.strip()
                    
                    if not text:
                        continue
                    
                    # 判断是否是标题
                    is_heading, level = self._is_heading(para)
                    
                    if is_heading:
                        # 创建新章节
                        section = {
                            "level": level,
                            "title": text,
                            "paragraphs": []
                        }
                        structure["sections"].append(section)
                        current_section = section
                        self.current_section = text
                    else:
                        # 普通段落
                        para_data = {
                            "index": paragraph_index,
                            "text": text,
                            "section": self.current_section,
                            "style": para.style.name if para.style else "Normal",
                            "char_count": len(text)
                        }
                        
                        structure["paragraphs"].append(para_data)
                        
                        if current_section:
                            current_section["paragraphs"].append(para_data)
                        
                        paragraph_index += 1
                        structure["metadata"]["total_chars"] += len(text)
                
                # 处理表格（可选）
                elif isinstance(element, CT_Tbl):
                    # 表格处理逻辑
                    pass
            
            structure["metadata"]["total_paragraphs"] = paragraph_index
            
            logger.info(f"文档解析完成: {paragraph_index} 个段落, {structure['metadata']['total_chars']} 个字符")
            
            return structure
        
        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            raise
    
    def _extract_title(self, doc: Document) -> str:
        """提取文档标题"""
        # 尝试从第一个段落提取
        if doc.paragraphs:
            first_para = doc.paragraphs[0]
            if first_para.style.name.startswith('Heading') or \
               first_para.style.name == 'Title':
                return first_para.text.strip()
        
        return "未命名文档"
    
    def _is_heading(self, para: Paragraph) -> tuple[bool, int]:
        """
        判断段落是否是标题
        
        Returns:
            (是否是标题, 标题级别)
        """
        style_name = para.style.name if para.style else ""
        
        # 检查样式名称
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading', '').strip())
                return True, level
            except:
                return True, 1
        
        # 检查文本模式（如：一、二、三 或 1. 2. 3.）
        text = para.text.strip()
        
        # 中文数字标题
        chinese_pattern = r'^[一二三四五六七八九十]+[、．]'
        if re.match(chinese_pattern, text):
            return True, 1
        
        # 阿拉伯数字标题
        number_pattern = r'^\d+[\.\．、]'
        if re.match(number_pattern, text):
            # 根据数字层级判断
            if re.match(r'^\d+[\.\．、]', text):
                return True, 1
            elif re.match(r'^\d+\.\d+[\.\．、]', text):
                return True, 2
        
        return False, 0
    
    def extract_text_only(self, file_path: str) -> str:
        """
        仅提取纯文本（用于简单场景）
        
        Args:
            file_path: 文档路径
        
        Returns:
            纯文本内容
        """
        try:
            doc = Document(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            return "\n".join(full_text)
        
        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            raise

