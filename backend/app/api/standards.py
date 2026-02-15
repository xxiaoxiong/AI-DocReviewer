"""
API 路由 - 标准文件管理接口
"""
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from typing import Optional, List
import json
import os
from pathlib import Path
from loguru import logger
import shutil

from ..tools.standard_converter import StandardConverter
from ..services.llm_service import LLMService
from ..core.rag_engine_v2 import RAGEngineV2

router = APIRouter(prefix="/api/standards", tags=["标准管理"])

# 获取项目根目录
backend_dir = Path(__file__).parent.parent.parent
project_root = backend_dir.parent
standards_dir = project_root / "standards" / "protocols"
raw_standards_dir = project_root / "standards" / "raw"

# 确保目录存在
standards_dir.mkdir(parents=True, exist_ok=True)
raw_standards_dir.mkdir(parents=True, exist_ok=True)


@router.post("/upload")
async def upload_standard(
    file: UploadFile = File(..., description="标准文档（Word格式）"),
    protocol_id: Optional[str] = Form(None, description="协议ID（可选，默认从文件名生成）"),
    protocol_name: Optional[str] = Form(None, description="协议名称（可选，默认从文档提取）"),
    use_llm: bool = Form(False, description="是否使用LLM辅助转换（更智能但较慢）")
):
    """
    上传标准文档并转换为JSON
    
    流程：
    1. 保存原始Word文档到 standards/raw/
    2. 使用转换器提取规则
    3. 生成JSON到 standards/protocols/
    4. 返回转换结果
    
    Args:
        file: Word文档
        protocol_id: 协议ID（可选）
        protocol_name: 协议名称（可选）
        use_llm: 是否使用LLM辅助（默认False，使用规则提取）
    
    Returns:
        转换结果
    """
    # 验证文件类型
    if not file.filename.endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="只支持 Word 文档（.docx, .doc）")
    
    logger.info("=" * 80)
    logger.info(f"📤 收到标准文档上传: {file.filename}")
    logger.info(f"   使用LLM: {use_llm}")
    logger.info("=" * 80)
    
    try:
        # 1. 保存原始文件
        raw_file_path = raw_standards_dir / file.filename
        with open(raw_file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        logger.info(f"✅ 原始文件已保存: {raw_file_path}")
        
        # 2. 生成协议ID（如果未提供）
        if not protocol_id:
            protocol_id = Path(file.filename).stem.upper().replace(" ", "_").replace("-", "_")
        
        logger.info(f"📋 协议ID: {protocol_id}")
        
        # 3. 转换文档
        output_path = standards_dir / f"{protocol_id}.json"
        
        if use_llm:
            # 使用LLM辅助转换（更智能）
            logger.info("🤖 使用LLM辅助转换...")
            converter = StandardConverter(use_llm=True)
            
            # 调用LLM服务进行智能提取
            result = await convert_with_llm(
                raw_file_path=str(raw_file_path),
                output_path=str(output_path),
                protocol_id=protocol_id,
                protocol_name=protocol_name
            )
        else:
            # 使用规则提取（快速）
            logger.info("⚡ 使用规则提取...")
            converter = StandardConverter(use_llm=False)
            converter.convert_word_to_json(
                word_path=str(raw_file_path),
                output_path=str(output_path),
                protocol_id=protocol_id,
                protocol_name=protocol_name
            )
            
            # 读取转换结果
            with open(output_path, 'r', encoding='utf-8') as f:
                result = json.load(f)
        
        # 4. 统计信息
        total_categories = len(result.get('categories', []))
        total_rules = sum(len(cat.get('rules', [])) for cat in result.get('categories', []))
        
        logger.info("=" * 80)
        logger.info(f"✅ 转换完成！")
        logger.info(f"   协议ID: {result.get('protocol_id')}")
        logger.info(f"   协议名称: {result.get('name')}")
        logger.info(f"   分类数: {total_categories}")
        logger.info(f"   规则数: {total_rules}")
        logger.info("=" * 80)
        
        # 自动重新加载RAG引擎
        try:
            from ..api import review
            review.rag_engine._load_standards()
            review.rag_engine._build_vector_index()
            logger.info("✅ RAG引擎已自动重新加载")
        except Exception as e:
            logger.warning(f"⚠️ RAG引擎重新加载失败: {e}")
        
        return {
            "success": True,
            "message": "标准文档转换成功",
            "data": {
                "protocol_id": result.get('protocol_id'),
                "protocol_name": result.get('name'),
                "version": result.get('version'),
                "total_categories": total_categories,
                "total_rules": total_rules,
                "file_path": str(output_path),
                "categories": [
                    {
                        "category": cat.get('category'),
                        "rule_count": len(cat.get('rules', []))
                    }
                    for cat in result.get('categories', [])
                ]
            }
        }
    
    except Exception as e:
        logger.error("=" * 80)
        logger.error(f"❌ 转换失败: {e}")
        logger.error("=" * 80)
        logger.exception("详细错误:")
        
        raise HTTPException(status_code=500, detail=f"转换失败: {str(e)}")


async def convert_with_llm(
    raw_file_path: str,
    output_path: str,
    protocol_id: str,
    protocol_name: Optional[str] = None
) -> dict:
    """
    使用LLM辅助转换标准文档
    
    Args:
        raw_file_path: 原始Word文档路径
        output_path: 输出JSON路径
        protocol_id: 协议ID
        protocol_name: 协议名称
    
    Returns:
        转换后的JSON数据
    """
    from docx import Document
    from ..models.document import Standard, Category, Rule, CheckType, Severity
    
    # 1. 读取Word文档
    doc = Document(raw_file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    
    # 2. 提取标题
    if not protocol_name:
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            if text and len(text) < 50:
                protocol_name = text
                break
        if not protocol_name:
            protocol_name = "未命名标准"
    
    logger.info(f"📄 文档内容长度: {len(full_text)} 字符")
    
    # 3. 使用LLM提取规则
    llm_service = LLMService(use_local=False)
    
    # 构造提取prompt
    prompt = f"""你是一个专业的标准文档分析助手。请从以下标准文档中提取规则。

【文档标题】
{protocol_name}

【文档内容】
{full_text[:8000]}  # 限制长度避免超token

【任务要求】
1. 识别文档中的章节（作为分类）
2. 提取每个章节下的规则条目
3. 分析每条规则的类型（format/semantic/structure）
4. 提取关键词
5. 判断严重程度（high/medium/low）

【输出格式】
严格按照以下JSON格式输出：
{{
  "categories": [
    {{
      "category": "分类名称",
      "rules": [
        {{
          "rule_id": "R001",
          "description": "规则描述",
          "check_type": "semantic",
          "keywords": ["关键词1", "关键词2"],
          "positive_examples": ["正确示例"],
          "negative_examples": ["错误示例"],
          "severity": "medium"
        }}
      ]
    }}
  ]
}}

【注意事项】
- check_type 只能是: semantic, format, structure
- severity 只能是: high, medium, low
- 尽量为每条规则生成1-2个正反例
- 关键词要精准，不超过5个

现在开始提取，只返回JSON，不要其他内容。
"""
    
    messages = [
        {
            "role": "system",
            "content": "你是标准文档分析专家，擅长从文档中提取结构化规则。你必须严格按照JSON格式输出。"
        },
        {
            "role": "user",
            "content": prompt
        }
    ]
    
    logger.info("🤖 调用LLM提取规则...")
    response = await llm_service.chat(
        messages=messages,
        temperature=0.1,
        max_tokens=4000,
        response_format={"type": "json_object"}
    )
    
    # 4. 解析LLM响应
    content = response["choices"][0]["message"]["content"]
    
    try:
        extracted_data = json.loads(content)
    except json.JSONDecodeError:
        logger.error(f"LLM返回的不是有效JSON: {content}")
        raise Exception("LLM返回格式错误")
    
    # 5. 构建标准对象
    categories = []
    for cat_data in extracted_data.get('categories', []):
        rules = []
        for rule_data in cat_data.get('rules', []):
            try:
                rule = Rule(
                    rule_id=rule_data.get('rule_id', 'R000'),
                    description=rule_data.get('description', ''),
                    check_type=CheckType(rule_data.get('check_type', 'semantic')),
                    keywords=rule_data.get('keywords', []),
                    positive_examples=rule_data.get('positive_examples', []),
                    negative_examples=rule_data.get('negative_examples', []),
                    severity=Severity(rule_data.get('severity', 'medium'))
                )
                rules.append(rule)
            except Exception as e:
                logger.warning(f"跳过无效规则: {e}")
        
        if rules:
            category = Category(
                category=cat_data.get('category', '未分类'),
                rules=rules
            )
            categories.append(category)
    
    standard = Standard(
        protocol_id=protocol_id,
        name=protocol_name,
        version="1.0",
        description=f"从 {Path(raw_file_path).name} 通过LLM提取",
        categories=categories
    )
    
    # 6. 保存JSON
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(
            standard.model_dump(),
            f,
            ensure_ascii=False,
            indent=2
        )
    
    logger.info(f"✅ LLM提取完成，共 {len(categories)} 个分类")
    
    return standard.model_dump()


@router.get("/list")
async def list_standards():
    """
    列出所有已转换的标准
    
    Returns:
        标准列表
    """
    try:
        standards = []
        
        for json_file in standards_dir.glob("*.json"):
            try:
                with open(json_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    
                    total_rules = sum(len(cat.get('rules', [])) for cat in data.get('categories', []))
                    
                    standards.append({
                        "protocol_id": data.get('protocol_id'),
                        "name": data.get('name'),
                        "version": data.get('version'),
                        "description": data.get('description'),
                        "total_categories": len(data.get('categories', [])),
                        "total_rules": total_rules,
                        "file_name": json_file.name
                    })
            except Exception as e:
                logger.warning(f"读取标准文件失败 {json_file}: {e}")
        
        return {
            "total": len(standards),
            "standards": standards
        }
    
    except Exception as e:
        logger.error(f"列出标准失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{protocol_id}")
async def get_standard_detail(protocol_id: str):
    """
    获取标准详情
    
    Args:
        protocol_id: 协议ID
    
    Returns:
        标准详细信息
    """
    try:
        json_file = standards_dir / f"{protocol_id}.json"
        
        if not json_file.exists():
            raise HTTPException(status_code=404, detail=f"标准 {protocol_id} 不存在")
        
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return data
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取标准详情失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/{protocol_id}")
async def delete_standard(protocol_id: str):
    """
    删除标准
    
    Args:
        protocol_id: 协议ID
    
    Returns:
        删除结果
    """
    try:
        json_file = standards_dir / f"{protocol_id}.json"
        
        if not json_file.exists():
            raise HTTPException(status_code=404, detail=f"标准 {protocol_id} 不存在")
        
        # 删除JSON文件
        json_file.unlink()
        
        # 尝试删除原始文件（如果存在）
        for raw_file in raw_standards_dir.glob(f"{protocol_id}.*"):
            raw_file.unlink()
            logger.info(f"已删除原始文件: {raw_file}")
        
        logger.info(f"✅ 已删除标准: {protocol_id}")
        
        # 自动重新加载RAG引擎
        try:
            from ..api import review
            review.rag_engine._load_standards()
            review.rag_engine._build_vector_index()
            logger.info("✅ RAG引擎已自动重新加载")
        except Exception as e:
            logger.warning(f"⚠️ RAG引擎重新加载失败: {e}")
        
        return {
            "success": True,
            "message": f"标准 {protocol_id} 已删除"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除标准失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{protocol_id}/download")
async def download_standard(protocol_id: str):
    """
    下载标准JSON文件
    
    Args:
        protocol_id: 协议ID
    
    Returns:
        JSON文件
    """
    try:
        json_file = standards_dir / f"{protocol_id}.json"
        
        if not json_file.exists():
            raise HTTPException(status_code=404, detail=f"标准 {protocol_id} 不存在")
        
        return FileResponse(
            path=str(json_file),
            filename=f"{protocol_id}.json",
            media_type="application/json"
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"下载标准失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/reload")
async def reload_all_standards():
    """
    重新加载所有标准到RAG引擎
    
    Returns:
        重载结果
    """
    try:
        # 重新初始化RAG引擎（会重新加载所有标准）
        from ..api import review
        
        logger.info("🔄 开始重新加载所有标准...")
        
        # 重新构建向量索引
        review.rag_engine._load_standards()
        review.rag_engine._build_vector_index()
        
        # 获取加载的标准数量
        total_standards = len(review.rag_engine.standards)
        total_rules = len(review.rag_engine.rule_index)
        
        logger.info(f"✅ 已重新加载 {total_standards} 个标准，共 {total_rules} 条规则")
        
        return {
            "success": True,
            "message": f"已重新加载 {total_standards} 个标准",
            "data": {
                "total_standards": total_standards,
                "total_rules": total_rules
            }
        }
    
    except Exception as e:
        logger.error(f"重新加载标准失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/batch-upload")
async def batch_upload_standards(
    files: List[UploadFile] = File(..., description="多个标准文档"),
    use_llm: bool = Form(False, description="是否使用LLM辅助")
):
    """
    批量上传标准文档
    
    Args:
        files: 多个Word文档
        use_llm: 是否使用LLM辅助
    
    Returns:
        批量转换结果
    """
    results = []
    
    for file in files:
        try:
            # 调用单个上传接口
            result = await upload_standard(
                file=file,
                protocol_id=None,
                protocol_name=None,
                use_llm=use_llm
            )
            results.append({
                "file_name": file.filename,
                "success": True,
                "data": result.get('data')
            })
        except Exception as e:
            logger.error(f"批量上传失败 {file.filename}: {e}")
            results.append({
                "file_name": file.filename,
                "success": False,
                "error": str(e)
            })
    
    success_count = sum(1 for r in results if r['success'])
    
    return {
        "total": len(files),
        "success": success_count,
        "failed": len(files) - success_count,
        "results": results
    }

